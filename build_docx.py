from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "informe_hmm.docx"
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)


def configurar_documento(documento):
    seccion = documento.sections[0]
    seccion.top_margin = Inches(1)
    seccion.bottom_margin = Inches(1)
    seccion.left_margin = Inches(1)
    seccion.right_margin = Inches(1)

    estilos = documento.styles
    estilos["Normal"].font.name = "Arial"
    estilos["Normal"].font.size = Pt(12)
    estilos["Title"].font.name = "Arial"
    estilos["Title"].font.size = Pt(22)
    estilos["Title"].font.bold = True
    estilos["Title"].font.color.rgb = ACCENT
    estilos["Heading 1"].font.name = "Arial"
    estilos["Heading 1"].font.size = Pt(16)
    estilos["Heading 1"].font.bold = True
    estilos["Heading 1"].font.color.rgb = ACCENT
    estilos["Heading 2"].font.name = "Arial"
    estilos["Heading 2"].font.size = Pt(14)
    estilos["Heading 2"].font.bold = True
    estilos["Heading 2"].font.color.rgb = ACCENT

    encabezado = seccion.header.paragraphs[0]
    encabezado.text = "Modelo Oculto de Markov"
    encabezado.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    encabezado.runs[0].font.name = "Arial"
    encabezado.runs[0].font.size = Pt(9)
    encabezado.runs[0].font.color.rgb = MUTED

    pie = seccion.footer.paragraphs[0]
    pie.text = "Actividad didactica 2-M2"
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pie.runs[0].font.name = "Arial"
    pie.runs[0].font.size = Pt(9)
    pie.runs[0].font.color.rgb = MUTED


def agregar_parrafo(documento, texto):
    parrafo = documento.add_paragraph(texto)
    parrafo.paragraph_format.space_after = Pt(6)
    parrafo.paragraph_format.line_spacing = 1.08
    for run in parrafo.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return parrafo


def agregar_titulo(documento, texto, nivel=1):
    parrafo = documento.add_heading(texto, level=nivel)
    for run in parrafo.runs:
        run.font.name = "Arial"
        run.font.color.rgb = ACCENT
    return parrafo


def agregar_vineta(documento, texto):
    parrafo = documento.add_paragraph(style="List Bullet")
    parrafo.paragraph_format.space_after = Pt(4)
    run = parrafo.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(12)


def agregar_numero(documento, texto):
    parrafo = documento.add_paragraph(style="List Number")
    parrafo.paragraph_format.space_after = Pt(4)
    run = parrafo.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(12)


def poner_texto_celda(celda, texto, negrita=False, alineacion=WD_ALIGN_PARAGRAPH.LEFT):
    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    parrafo = celda.paragraphs[0]
    parrafo.alignment = alineacion
    run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
    run.text = texto
    run.bold = negrita
    run.font.name = "Arial"
    run.font.size = Pt(10)


def agregar_tabla(documento, encabezados, filas):
    tabla = documento.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"

    for indice, encabezado in enumerate(encabezados):
        poner_texto_celda(tabla.rows[0].cells[indice], encabezado, True, WD_ALIGN_PARAGRAPH.CENTER)

    for fila in filas:
        celdas = tabla.add_row().cells
        for indice, valor in enumerate(fila):
            alineacion = WD_ALIGN_PARAGRAPH.LEFT if indice == 0 else WD_ALIGN_PARAGRAPH.CENTER
            poner_texto_celda(celdas[indice], str(valor), False, alineacion)

    documento.add_paragraph()


def construir_documento():
    documento = Document()
    configurar_documento(documento)

    titulo = documento.add_paragraph(style="Title")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.add_run("Implementacion de un Modelo Oculto de Markov aplicado al clima")

    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Estados ocultos, observaciones visibles y simulacion probabilistica")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED
    documento.add_paragraph()

    agregar_titulo(documento, "Introduccion")
    agregar_parrafo(
        documento,
        "Un Modelo Oculto de Markov, o HMM, representa procesos donde existe una parte que no se observa directamente. "
        "Esa parte se llama estado oculto. Aunque el estado oculto no se ve, el sistema produce observaciones visibles que funcionan como pistas.",
    )
    agregar_parrafo(
        documento,
        "En esta actividad se simulan condiciones atmosfericas no observables directamente y se relacionan con resultados visibles del clima. "
        "Los estados ocultos son Alta presion y Baja presion. Las observaciones visibles son Soleado, Nublado y Lluvioso.",
    )

    agregar_vineta(documento, "Probabilidades iniciales: indican en que estado oculto puede comenzar el sistema.")
    agregar_vineta(documento, "Probabilidades de transicion: indican como cambia el sistema entre estados ocultos.")
    agregar_vineta(documento, "Probabilidades de emision: indican que observacion visible puede aparecer desde cada estado oculto.")

    agregar_titulo(documento, "Metodologia")
    agregar_parrafo(
        documento,
        "La solucion se implemento en Python sin librerias externas. El fenomeno elegido fue el clima diario, donde la presion atmosferica actua como causa interna y el clima visible actua como evidencia observable.",
    )

    agregar_titulo(documento, "Estados y observaciones", 2)
    agregar_tabla(
        documento,
        ["Tipo", "Valores"],
        [
            ["Estados ocultos", "Alta presion, Baja presion"],
            ["Observaciones visibles", "Soleado, Nublado, Lluvioso"],
        ],
    )

    agregar_titulo(documento, "Probabilidades iniciales", 2)
    agregar_tabla(
        documento,
        ["Estado oculto", "Probabilidad inicial"],
        [["Alta presion", "0.60"], ["Baja presion", "0.40"]],
    )

    agregar_titulo(documento, "Matriz de transicion", 2)
    agregar_tabla(
        documento,
        ["Estado actual", "Alta presion", "Baja presion"],
        [["Alta presion", "0.75", "0.25"], ["Baja presion", "0.35", "0.65"]],
    )

    agregar_titulo(documento, "Matriz de emision", 2)
    agregar_tabla(
        documento,
        ["Estado oculto", "Soleado", "Nublado", "Lluvioso"],
        [["Alta presion", "0.70", "0.25", "0.05"], ["Baja presion", "0.15", "0.35", "0.50"]],
    )

    agregar_titulo(documento, "Algoritmo implementado", 2)
    for paso in [
        "Validar probabilidades iniciales, matriz de transicion y matriz de emision.",
        "Elegir el estado oculto inicial.",
        "Generar una observacion visible desde ese estado oculto.",
        "Cambiar al siguiente estado oculto usando la matriz de transicion.",
        "Generar una nueva observacion usando la matriz de emision.",
        "Repetir el proceso y ejecutar muchas simulaciones para obtener resultados agregados.",
    ]:
        agregar_numero(documento, paso)

    documento.add_section(WD_SECTION.NEW_PAGE)
    agregar_titulo(documento, "Resultados y conclusiones")
    agregar_parrafo(
        documento,
        "Se ejecuto una prueba con 1000 simulaciones, 20 pasos por simulacion y semilla 42. La simulacion produjo secuencias de estados ocultos y observaciones visibles.",
    )

    agregar_titulo(documento, "Resultados agregados", 2)
    agregar_tabla(
        documento,
        ["Medida", "Alta presion", "Baja presion"],
        [
            ["Distribucion observada de estados ocultos", "0.5806", "0.4194"],
            ["Distribucion estacionaria aproximada", "0.5833", "0.4167"],
        ],
    )
    agregar_tabla(
        documento,
        ["Observacion visible", "Frecuencia observada"],
        [["Soleado", "0.4677"], ["Nublado", "0.2959"], ["Lluvioso", "0.2364"]],
    )

    agregar_parrafo(
        documento,
        "Los resultados muestran que el sistema tiende a pasar mas tiempo en Alta presion. Tambien se observa que Soleado aparece con mayor frecuencia, lo cual es coherente con la matriz de emision.",
    )
    agregar_parrafo(
        documento,
        "Como conclusion, el HMM permite representar procesos donde existe informacion no observable directamente. En este caso, las condiciones de presion atmosferica son los estados ocultos y el clima visible son las observaciones.",
    )

    documento.save(OUTPUT)


if __name__ == "__main__":
    construir_documento()
